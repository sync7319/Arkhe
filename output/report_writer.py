"""
Writes the executive report as a professional Word document (.docx).
Parses the structured markdown output from report_agent and maps it to Word styles.
"""
import os
import re
from datetime import date
from config.settings import OUTPUT_DIR


def write_report(report_text: str, repo_path: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    out_dir  = os.path.join(repo_path, OUTPUT_DIR)
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "EXECUTIVE_REPORT.docx")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Cover block ───────────────────────────────────────────
    title_para = doc.add_heading("Executive Codebase Report", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    repo_name = os.path.basename(os.path.abspath(repo_path))
    sub = doc.add_paragraph(f"{repo_name}  ·  {date.today().isoformat()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # ── Body ──────────────────────────────────────────────────
    _render_body(doc, report_text)

    doc.save(out_path)
    return out_path


def _render_body(doc, text: str) -> None:
    """Convert the structured markdown report into Word paragraphs."""
    for line in text.splitlines():
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)

        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)

        elif stripped.startswith(("- ", "* ", "• ")):
            content = stripped[2:].strip()
            doc.add_paragraph(_strip_inline_md(content), style="List Bullet")

        elif re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s*", "", stripped)
            doc.add_paragraph(_strip_inline_md(content), style="List Number")

        else:
            doc.add_paragraph(_strip_inline_md(stripped))


def _strip_inline_md(text: str) -> str:
    """Remove bold/italic/code markdown markers for clean Word output."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    return text
